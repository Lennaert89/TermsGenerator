import os
import json
import csv
import docx
import argparse
import logging
from docx import Document
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import markdown

def setup_logger(log_file=None, verbosity='INFO'):
    log_format = '%(asctime)s - %(levelname)s - %(message)s'
    logger = logging.getLogger()
    logger.setLevel(getattr(logging, verbosity.upper(), 'INFO'))

    console_handler = logging.StreamHandler()
    console_handler.setFormatter(logging.Formatter(log_format))
    logger.addHandler(console_handler)

    if log_file:
        file_handler = logging.FileHandler(log_file)
        file_handler.setFormatter(logging.Formatter(log_format))
        logger.addHandler(file_handler)

def read_input_file(file_path):
    logging.info(f"Reading input file: {file_path}")
    _, file_extension = os.path.splitext(file_path)
    try:
        if file_extension.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_extension.lower() == '.pdf':
            pdf_reader = PdfReader(file_path)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        elif file_extension.lower() == '.docx':
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_extension.lower() == '.md':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        logging.error(f"Error reading input file: {e}")
        raise

def load_dictionaries(dict_paths):
    terms_dict = {}
    dict_files = []

    for dict_path in dict_paths:
        if os.path.isdir(dict_path):
            for root, _, files in os.walk(dict_path):
                for file in files:
                    if file.endswith('.json') or file.endswith('.csv'):
                        dict_files.append(os.path.join(root, file))
        else:
            dict_files.append(dict_path)

    for dict_file in dict_files:
        logging.info(f"Loading dictionary file: {dict_file}")
        _, file_extension = os.path.splitext(dict_file)
        try:
            if file_extension.lower() == '.json':
                with open(dict_file, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    for entry in data:
                        terms_dict[entry['word'].lower()] = {
                            "meaning": entry['meaning'],
                            "reference": entry.get('reference', '')
                        }
            elif file_extension.lower() == '.csv':
                with open(dict_file, 'r', encoding='utf-8') as file:
                    reader = csv.DictReader(file)
                    for row in reader:
                        terms_dict[row['word'].lower()] = {
                            "meaning": row['meaning'],
                            "reference": row.get('reference', '')
                        }
            else:
                raise ValueError(f"Unsupported dictionary format: {file_extension}")
        except Exception as e:
            logging.error(f"Error loading dictionary file: {e}")
            raise
    return terms_dict

def find_terms_in_text(text, terms_dict):
    logging.info("Searching for terms in the text.")
    found_terms = {}
    text_lower = text.lower()
    for term, details in terms_dict.items():
        if term in text_lower:
            found_terms[term] = details
    logging.info(f"Found {len(found_terms)} terms in the text.")
    return found_terms

def save_to_docx(terms, output_path, also_see_text):
    logging.info(f"Saving output to DOCX file: {output_path}")
    try:
        doc = Document()
        for term, details in terms.items():
            doc.add_heading(term, level=1)
            doc.add_paragraph(details['meaning'])
            if details['reference']:
                doc.add_paragraph(f"{also_see_text} {details['reference']}", style='Italic')
        doc.save(output_path)
    except Exception as e:
        logging.error(f"Error saving to DOCX file: {e}")
        raise

def save_to_html(terms, output_path, also_see_text):
    logging.info(f"Saving output to HTML file: {output_path}")
    try:
        soup = BeautifulSoup("<html><body></body></html>", 'html.parser')
        body = soup.body
        for term, details in terms.items():
            term_header = soup.new_tag('h1')
            term_header.string = term
            body.append(term_header)
            if details['meaning'] is not None:
                meaning_paragraph = soup.new_tag('p')
                meaning_paragraph.string = details['meaning']
                body.append(meaning_paragraph)
            if details['reference']:
                reference_paragraph = soup.new_tag('p', style='font-style:italic;')
                reference_paragraph.string = f"{also_see_text} {details['reference']}"
                body.append(reference_paragraph)
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(soup.prettify())
    except Exception as e:
        logging.error(f"Error saving to HTML file: {e}")
        raise

def save_to_txt(terms, output_path, also_see_text):
    logging.info(f"Saving output to TXT file: {output_path}")
    try:
        with open(output_path, 'w', encoding='utf-8') as file:
            for term, details in terms.items():
                file.write(f"{term}\n{details['meaning']}\n")
                if details['reference']:
                    file.write(f"{also_see_text} {details['reference']}\n")
                file.write("\n")
    except Exception as e:
        logging.error(f"Error saving to TXT file: {e}")
        raise

def save_to_md(terms, output_path, also_see_text):
    logging.info(f"Saving output to Markdown file: {output_path}")
    try:
        with open(output_path, 'w', encoding='utf-8') as file:
            for term, details in terms.items():
                file.write(f"# {term}\n\n{details['meaning']}\n\n")
                if details['reference']:
                    file.write(f"*{also_see_text} {details['reference']}*\n\n")
    except Exception as e:
        logging.error(f"Error saving to Markdown file: {e}")
        raise

def main(input_file, dict_paths, output_format='docx', output_file=None, log_file=None, verbosity='INFO', language='en'):
    setup_logger(log_file, verbosity)

    if language == 'nl':
        also_see_text = "Zie ook:"
    else:
        also_see_text = "Also see:"

    text = read_input_file(input_file)
    terms_dict = load_dictionaries(dict_paths)
    found_terms = find_terms_in_text(text, terms_dict)

    if output_file is None:
        output_file = f"output.{output_format}"

    if output_format == 'docx':
        save_to_docx(found_terms, output_file, also_see_text)
    elif output_format == 'html':
        save_to_html(found_terms, output_file, also_see_text)
    elif output_format == 'txt':
        save_to_txt(found_terms, output_file, also_see_text)
    elif output_format == 'md':
        save_to_md(found_terms, output_file, also_see_text)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract terms and definitions from a file.")
    parser.add_argument('input_file', help="The input file to scan (docx, pdf, txt, md).")
    parser.add_argument('dict_paths', nargs='+', help="The dictionary files or directories to use (csv, json).")
    parser.add_argument('--output_format', choices=['docx', 'html', 'txt', 'md'], default='docx', help="The output format (docx, html, txt, md). Default is docx.")
    parser.add_argument('--output_file', help="The output file name. If not specified, defaults to 'output.<format>'.")
    parser.add_argument('--log', help="The log file to write to. If not specified, logs will only be printed to the terminal.")
    parser.add_argument('--verbosity', choices=['DEBUG', 'INFO', 'WARNING', 'ERROR', 'CRITICAL'], default='INFO', help="The verbosity level of logging. Default is INFO.")
    parser.add_argument('--language', choices=['en', 'nl'], default='en', help="The language for the 'Also see' text. Default is English.")
    
    args = parser.parse_args()
    main(args.input_file, args.dict_paths, args.output_format, args.output_file, args.log, args.verbosity, args.language)
